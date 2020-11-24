from docx import Document
from docx.shared import RGBColor, Pt, Cm
import os
# import glob
import pycorrector

def printCopyRight():
    print("-------------------------------文本AI分析---------------------------------------")
    print("******************************************************************************")
    print("****用途：此软件著作权权为个人所有，用于人工智能研究与学习，未经授权不得用于商业用途***********")
    print("****说明：此软件使用了百度智能平台的自然语言处理服务，在免费使用额度使用本程序不会产生费用******")
    print("*如果因为使用者需要使用超过免费额度的服务而产生的费用，与软件作者无关。**********************")
    print("-------------------------------------------------------------------------------")

def printMenu():
    print("\n-------------------------------------------------------------------------------")
    print("请选择：")
    print("1、输入docx格式的带完整路径的文件名")
    print("2、输入txt格式的带完整路径的文件名")
    print("3、退出")
    print("----------------------------------------------------------------------------------\n")

def readText(filename):
    textArray = []
    with open(filename, "r", errors='ignore') as f1:
        lines = f1.readlines()
        for line in lines:
            textArray.append(line.strip())
    return textArray

def readDoc(filename):
    doc = Document(filename)
    #store the paragraph of the docx
    paraArray = []
    for paragraph in doc.paragraphs:
        paraArray.append(paragraph.text.strip())
    return paraArray

def getNewFileName(filename):
    nameSplit = filename.split('.')
    newName = nameSplit[0]
    return newName + "错别字分析" + '.docx'

def flagDocx(filename):
    try:
        doc = Document(filename)
        allErrWord = []
        errCount = 0
        print("开始进行错别字分析......")
        with open('d:\\错别字分析.doc', 'w') as f:
            for paragraph in doc.paragraphs:
                # print("{}{}".format(i,paragraph.text))
                for run in paragraph.runs:
                    # idx_errors = pycorrector.detect(run.text)
                    corrected_sent, detail = pycorrector.correct(run.text)
                    if len(detail) != 0:
                        #three are error words, flag them in the docx
                        for elem in detail:
                            errCount += 1
                            print("\n---------------------------------------------------------------")
                            print("{}\n这段文本存在错别字".format(run.text))
                            print("发现错别字:{},正确词为:{}".format(elem[0],elem[1]))
                            print("错别字位置在本段的第{}到第{}字之间".format(elem[2],elem[3]))
                            print("---------------------------------------------------------------\n")
                            f.write("\n---------------------------------------------------------------")
                            f.write("{}\n这段文本存在错别字".format(run.text))
                            f.write("发现错别字:{},正确词为:{}".format(elem[0],elem[1]))
                            f.write("---------------------------------------------------------------\n")
                            allErrWord.append(elem[0])
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 0, 0)
                        print("\n*********************************************************************")
                        print("正确的句子是:\n{}".format(corrected_sent))
                        print("*********************************************************************\n")
            f.write("\n------------------------------------------------------------------------------\n")
            f.write("总共发现{}处错别字".format(errCount))
            f.write("\n------------------------------------------------------------------------------\n")
        print("正在将错别字进行标识，存到文件{}中...".format(getNewFileName(filename)))
        doc.save(getNewFileName(filename))
        print("----------------------------------------------------------------------------------")
        print("总共发现{}处错别字,所有错别字为:{}".format(errCount,allErrWord))
        print("错别字所在句子已经用红色粗体标识出，存储在 {}文件中".format(getNewFileName(filename)))
        print("-----------------------------------------------------------------------------------")
        return allErrWord
    except Exception as e:
        print(e)
        return None


#analysis the text,if there are wrong words, return the Correct text.Otherwise return None
def textAnalysis(text):
    errs = 0
    try:
        corrected_sent, detail = pycorrector.correct(text)
        if len(detail) != 0:
            # three are error words, flag them in the docx
            for elem in detail:
                errs += 1
                print("---------------------------------------------------------------")
                print("{}\n这段文本存在错别字".format(text))
                print("发现错别字:{},正确词为:{}".format(elem[0], elem[1]))
                print("错别字位置在本段的第{}到第{}字之间".format(elem[2], elem[3]))
                print("---------------------------------------------------------------")
            return (corrected_sent,errs)
        else:
            return (corrected_sent,0)
    except:
        return (None,0)
"""
    parse the txt file and read one line each time.Analysis each line 
    whether existed wrong words
"""
def parsTxtFile(filename):
    lineNo = 0
    # count the line number which has wrong words
    errCount = 0
    try:
        with open(filename, "r", errors='ignore') as f1:  # add encoding='UTF-8' or change 'r' to 'rb' or errors='ignore'
            print('开始读原始文件,进行错别字分析......')
            lines = f1.readlines()
            with open('d:\\错别字分析.doc', 'w') as f2:
                for line in lines:
                    lineNo += 1
                    textCorrected,errs = textAnalysis(line)
                    if textCorrected != None:
                        errCount += errs
                        # write the correct text into the file
                        f2.write("行号：{}有错别字，正确的句子是：".format(lineNo))
                        f2.write(textCorrected)
                    else:
                        print("line {}:{} OK".format(lineNo, line))
                f2.write("---------------------------------------------------\n")
                f2.write("总共完成{}行文本分析\n".format(lineNo))
                f2.write("总共有{}c处错别字\n".format(errCount))
                f2.write("---------------------------------------------------\n")
            print('*********************************************************')
            print('已经完成文本错别字分析，输出结果存储在D:\\错别字分析.doc文件中')
            print("总共完成{}行文本分析\n".format(lineNo))
            print("总共有{}处错别字\n".format(errCount))
            print('*********************************************************')
    except IOError as e:
        print(e)



def main():
    printCopyRight()
    printMenu()
    choice = ''
    while (choice != '3'):
        choice = input("请输入数字1/2/3进行选择: ")
        if choice == '1':
            filename = input("请输入待分析的docx的文件名: ")
            flagDocx(filename)
            printMenu()
        elif choice == '2':
            # to be developed
            # print("此功能开发中......")
            filename = input("请输入待分析的txt的文件名: ")
            parsTxtFile(filename)
            printMenu()

main()